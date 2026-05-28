"""
handlers/contracts.py — генерация договоров для менеджеров West Asia.

Флоу:
1. Менеджер пишет про договор → бот показывает 4 кнопки
2. Менеджер выбирает тип → бот просит карточку компании
3. Менеджер отправляет файл/текст → AI извлекает реквизиты
4. Бот заполняет шаблон .docx и отправляет готовый договор
"""
import re
import io
import os
import json
import asyncio
import tempfile
import logging
import requests
from datetime import datetime
from aiogram import Router, F
from aiogram.types import Message, ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove

from config import ADMIN_ID, logger
from services.ai import ask_deepseek

router = Router()

# ─── Конфигурация договоров ───────────────────────────────────────────────────
CONTRACTS = {
    'AraliaPostavka': {
        'name': 'Договор поставки Аралия Трек',
        'button': '📦 Договор поставки Аралия Трек',
        'template_url': 'https://drive.google.com/uc?export=download&id=1HNlzwXMgdu87yZ_BXnK1B1H6VQeo-nW_',
    },
    'AsiaPostavka': {
        'name': 'Договор поставки АЗИЯ ИМПОРТ',
        'button': '📦📦 Договор поставки АЗИЯ ИМПОРТ',
        'template_url': 'https://drive.google.com/uc?export=download&id=1HWspZL44BFmZVmO0ul-LSOOm8kKZVnAV',
    },
    'AraliaTEO': {
        'name': 'Договор ТЭО Аралия Трек',
        'button': '🚚 Договор ТЭО Аралия Трек',
        'template_url': 'https://drive.google.com/uc?export=download&id=1YMlcKayE61Jz5fH9APQ35PF1OMzDxwZs',
    },
    'AsiaTEO': {
        'name': 'Договор ТЭО АЗИЯ ИМПОРТ',
        'button': '🚛 Договор ТЭО АЗИЯ ИМПОРТ',
        'template_url': 'https://drive.google.com/uc?export=download&id=1EQix29pyUYDG0GyvaxJIb_Oq24ejOhlf',
    },
}

# ─── Данные менеджеров ────────────────────────────────────────────────────────
MANAGERS = {
    745001459:  {'prefix': 'Л1',  'name': 'Татьяна',    'position': 'Логист', 'phone': '+7 (977) 896-59-66', 'email': 'logist1@west-asia.com'},
    6005416734: {'prefix': 'Л2',  'name': 'Юлия',       'position': 'Логист', 'phone': '+7 (985) 510-03-35', 'email': 'logist2@west-asia.com'},
    7141105878: {'prefix': 'Л3',  'name': 'Роман',      'position': 'Логист', 'phone': '+7 (985) 857-25-53', 'email': 'logist3@west-asia.com'},
    980194477:  {'prefix': 'Л4',  'name': 'Мария',      'position': 'Логист', 'phone': '+7 (999) 558-08-84', 'email': 'logist4@west-asia.com'},
    8220688531: {'prefix': 'Л6',  'name': 'Мария',      'position': 'Логист', 'phone': '+7 (985) 211-75-13', 'email': 'logist6@west-asia.com'},
    853238336:  {'prefix': 'АБ7', 'name': 'Андрей Б.',  'position': 'Логист', 'phone': '+7 (985) 233-55-83', 'email': 'ab@west-asia.com'},
}
DEFAULT_MANAGER = {'prefix': '', 'name': '', 'position': '', 'phone': '', 'email': ''}

# Состояния: user_id → {'step': 'select'|'wait_card', 'contract_type': str}
CONTRACT_STATE: dict = {}

# Кеш шаблонов
_template_cache: dict = {}

# Счётчик номеров договоров
_counter_lock = asyncio.Lock()
_COUNTER_FILE = 'contract_counter.json'

# ─── Триггеры запуска флоу ───────────────────────────────────────────────────
CONTRACT_TRIGGERS = [
    "нужен договор", "составь договор", "сделай договор", "создай договор",
    "оформи договор", "договор поставки", "договор тэо", "договор с клиентом",
    "нужен контракт", "составить договор", "подготовь договор",
    "договор аралия", "договор азия", "новый договор",
    "надо договор", "нужно договор", "сделать договор",
    "хочу договор", "договор клиент", "подготовить договор",
    "договор",  # просто слово "договор" тоже триггер
]


def is_contract_request(text: str) -> bool:
    """Определяет является ли запрос про договор."""
    t = text.lower().strip()
    return any(trigger in t for trigger in CONTRACT_TRIGGERS)


# ─── Команда /contract ────────────────────────────────────────────────────────
from aiogram.filters import Command

@router.message(Command("contract"))
async def cmd_contract(message: Message):
    """Команда /contract — запуск флоу договоров."""
    await start_contract_flow(message)


# ─── Утилиты ─────────────────────────────────────────────────────────────────
def get_manager(user_id: int) -> dict:
    return MANAGERS.get(user_id, DEFAULT_MANAGER)


def get_contract_keyboard() -> ReplyKeyboardMarkup:
    buttons = [[KeyboardButton(text=c['button'])] for c in CONTRACTS.values()]
    buttons.append([KeyboardButton(text="❌ Отмена")])
    return ReplyKeyboardMarkup(keyboard=buttons, resize_keyboard=True, one_time_keyboard=True)


def get_current_date() -> dict:
    now = datetime.now()
    months = {1:'января',2:'февраля',3:'марта',4:'апреля',5:'мая',6:'июня',
              7:'июля',8:'августа',9:'сентября',10:'октября',11:'ноября',12:'декабря'}
    return {
        'date_full': f'«{now.day}» {months[now.month]} {now.year} г.',
        'date_day': now.day, 'date_month': months[now.month], 'date_year': now.year,
        'date_digital': now.strftime('%d.%m.%Y'),
        'date_digital_line': now.strftime('%d/%m/%Y'),
        'date_iso': now.strftime('%Y-%m-%d'),
    }


def get_gender(name: str) -> str:
    """Определяет пол по имени/отчеству. Возвращает 'f' или 'm'."""
    if not name:
        return 'm'
    parts = name.strip().split()
    # Проверяем отчество (3-е слово) или имя (2-е слово)
    check = parts[2] if len(parts) >= 3 else (parts[1] if len(parts) >= 2 else parts[0])
    check = check.rstrip('.')
    female_endings = ['вна', 'овна', 'евна', 'ична', 'инична', 'ьевна',
                      'а', 'я', 'ва', 'на', 'ина', 'ская', 'цкая']
    for e in female_endings:
        if check.lower().endswith(e):
            return 'f'
    return 'm'


def get_acting_form(ceo_name: str) -> str:
    """Возвращает 'действующего' или 'действующей' по имени."""
    return 'действующей' if get_gender(ceo_name) == 'f' else 'действующего'


def add_quotes(name: str) -> str:
    if not name or name == '___________':
        return name
    # Если кавычки уже есть — не трогаем
    if '«' in name or '»' in name:
        return name
    if re.match(r'^ИП\b', name, re.IGNORECASE):
        return name.strip()
    for pattern, repl in [
        (r'(ООО|АО|ПАО|ЗАО)\s+(.+)$', r'\1 «\2»'),
        (r'(Общество с ограниченной ответственностью)\s+(.+)$', r'\1 «\2»'),
    ]:
        if re.search(pattern, name, re.IGNORECASE):
            return re.sub(pattern, repl, name)
    return name


async def get_contract_number(prefix: str, contract_key: str) -> str:
    if not prefix:
        return ""
    now = datetime.now()
    year = now.strftime('%y')
    day_month = now.strftime('%d%m')
    base = f"{prefix}/{year}/{day_month}"
    async with _counter_lock:
        counters = {}
        if os.path.exists(_COUNTER_FILE):
            with open(_COUNTER_FILE, 'r', encoding='utf-8') as f:
                counters = json.load(f)
        key = f"{prefix}_{contract_key}_{day_month}"
        if key not in counters:
            counters[key] = 2
            with open(_COUNTER_FILE, 'w', encoding='utf-8') as f:
                json.dump(counters, f, ensure_ascii=False, indent=2)
            return base
        num = counters[key]
        counters[key] = num + 1
        with open(_COUNTER_FILE, 'w', encoding='utf-8') as f:
            json.dump(counters, f, ensure_ascii=False, indent=2)
        return f"{base}-{num}"


# ─── Извлечение текста из файлов ─────────────────────────────────────────────
def extract_text_docx(path: str) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return '\n'.join(lines)
    except Exception as e:
        logger.error(f"DOCX extract error: {e}")
        return ""


def extract_text_pdf(path: str) -> str:
    text = ''
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + '\n'
    except Exception:
        try:
            import PyPDF2
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ''
        except Exception:
            pass
    return text


def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == '.docx':
        return extract_text_docx(path)
    elif ext == '.pdf':
        return extract_text_pdf(path)
    elif ext in ('.xlsx', '.xls'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(c) for c in row if c]
                    if row_text:
                        lines.append(' | '.join(row_text))
            return '\n'.join(lines)
        except Exception as e:
            logger.error(f"XLSX extract error: {e}")
            return ""
    return ""


# ─── AI извлечение реквизитов ────────────────────────────────────────────────
async def extract_company_data(raw_text: str) -> dict:
    prompt = f"""Ты — эксперт по извлечению данных из карточек компаний и ИП. Верни ТОЛЬКО JSON.

ТЕКСТ КАРТОЧКИ:
{raw_text[:6000]}

ТРЕБУЕМЫЕ ПОЛЯ:
{{
    "company_name": "Полное название в именительном падеже. СОХРАНЯЙ ВСЕ КАВЫЧКИ как в оригинале, в том числе внутренние.",
    "company_name_short": "Короткое название. СОХРАНЯЙ ВСЕ КАВЫЧКИ как в оригинале.",
    "company_name_genitive": "Название в родительном падеже. СОХРАНЯЙ ВСЕ КАВЫЧКИ.",
    "inn": "ИНН - 10 или 12 цифр",
    "kpp": "КПП - 9 цифр (только для ООО, для ИП пустая строка)",
    "ogrn": "ОГРН - 13 цифр (только для ООО)",
    "ogrnip": "ОГРНИП - 15 цифр (только для ИП)",
    "ceo_name_nominative": "ФИО руководителя полностью в именительном падеже",
    "ceo_name_nominative_short": "ФИО руководителя сокращённо",
    "ceo_name_genitive": "ФИО руководителя в родительном падеже полностью",
    "ceo_name_genitive_short": "ФИО руководителя в родительном падеже сокращённо",
    "ceo_position_nominative": "Должность в именительном падеже. Для ООО: 'Генеральный директор'. Для ИП: 'Индивидуальный предприниматель'",
    "ceo_position_genitive": "Должность в родительном падеже. Для ООО: 'Генерального директора'. Для ИП: 'Индивидуального предпринимателя'",
    "is_individual_entrepreneur": "true если ИП, false если ООО",
    "legal_basis": "Основание. Для ООО: 'Устава'. Для ИП: 'ОГРНИП'. Если есть доверенность: 'Доверенности'",
    "legal_basis_full": "Полное основание. Для ООО: 'на основании Устава'. Для ИП: 'на основании ОГРНИП'. Если есть доверенность: сохраняй номер и дату.",
    "address": "Юридический адрес полностью",
    "phone": "Номер телефона",
    "email": "Email адрес",
    "bik": "БИК - 9 цифр",
    "account": "Расчётный счёт - 20 цифр",
    "korr_account": "Корреспондентский счёт - 20 цифр",
    "bank_name": "Название банка",
    "has_managing_company": "true/false — указана ли управляющая организация (УК)",
    "managing_company_name": "Полное название УК в именительном падеже, если есть",
    "managing_company_name_short": "Короткое название УК, если есть",
    "managing_ceo_name_nominative": "ФИО руководителя УК в именительном падеже, если есть",
    "managing_ceo_name_genitive": "ФИО руководителя УК в родительном падеже, если есть",
    "managing_ceo_name_nominative_short": "ФИО руководителя УК сокращённо, если есть",
    "managing_ceo_name_genitive_short": "ФИО руководителя УК в родительном падеже сокращённо, если есть",
    "managing_ceo_position_nominative": "Должность руководителя УК, если есть"
}}

ПРАВИЛА:
1. Если в тексте есть 'Индивидуальный предприниматель' или 'ИП' или ОГРНИП (15 цифр) — is_individual_entrepreneur = true
2. Для ИП: legal_basis = 'ОГРНИП', legal_basis_full = 'на основании ОГРНИП'
3. Для ООО: legal_basis = 'Устава', legal_basis_full = 'на основании Устава'
4. СОХРАНЯЙ КАВЫЧКИ в названиях компаний, в том числе внутренние!
5. Если в тексте есть 'доверенность': legal_basis = 'Доверенности', legal_basis_full = указанное в тексте
6. Если в тексте есть 'Управляющая организация' или 'управляющая организация':
   - has_managing_company = true
   - ceo_name_nominative = ФИО руководителя УК
   - ceo_position_nominative = 'Генеральный директор [короткое название УК] — управляющей организации [короткое название компании]'
   - ceo_position_genitive = 'Генерального директора [короткое название УК] — управляющей организации [короткое название компании]'
   - legal_basis_full = 'на основании Устава [короткое название компании] и Договора о передаче полномочий единоличного исполнительного органа от «___» ______ 20__ г. № ___'

Верни ТОЛЬКО JSON!"""

    try:
        answer = await ask_deepseek(
            messages=[
                {"role": "system", "content": "Возвращай только валидный JSON. Сохраняй кавычки в названиях компаний."},
                {"role": "user", "content": prompt}
            ]
        )
        # Очищаем от markdown
        answer = re.sub(r'^```json\s*', '', answer.strip())
        answer = re.sub(r'^```\s*', '', answer)
        answer = re.sub(r'\s*```$', '', answer)
        data = json.loads(answer)

        # Пост-обработка управляющей организации
        if data.get('has_managing_company'):
            mc_full = data.get('managing_company_name', '')
            mc_short = data.get('managing_company_name_short', '')
            # Фолбэк: если короткое название не извлечено — берём из полного
            if not mc_short and mc_full:
                m = re.search(r'(ООО|АО|ПАО|ЗАО)\s*[«"]?([^»"]+)[»"]?', mc_full)
                if m:
                    mc_short = f"{m.group(1)} «{m.group(2).strip()}»"
                    data['managing_company_name_short'] = mc_short
            co_short = data.get('company_name_short', '')
            if mc_short and co_short:
                data['ceo_position_nominative'] = f"Генеральный директор {mc_short} — управляющей организации {co_short}"
                data['ceo_position_genitive'] = f"Генерального директора {mc_short} — управляющей организации {co_short}"
            if 'Договора о передаче' not in data.get('legal_basis_full', ''):
                co_name = data.get('company_name_short', co_short)
                data['legal_basis_full'] = f"на основании Устава {co_name} и Договора о передаче полномочий единоличного исполнительного органа от «___» ______ 20__ г. № ___"

        # ИП обработка
        if data.get('ogrnip') or data.get('is_individual_entrepreneur'):
            data['is_individual_entrepreneur'] = True
            if not data.get('ceo_position_nominative'):
                data['ceo_position_nominative'] = 'Индивидуальный предприниматель'
                data['ceo_position_genitive'] = 'Индивидуального предпринимателя'
            data['legal_basis'] = 'ОГРНИП'
            data['legal_basis_full'] = 'на основании ОГРНИП'
        else:
            data['is_individual_entrepreneur'] = False
            if not data.get('ceo_position_nominative'):
                data['ceo_position_nominative'] = 'Генеральный директор'
                data['ceo_position_genitive'] = 'Генерального директора'

        # Кавычки в названиях
        for field in ('company_name', 'company_name_short', 'company_name_genitive'):
            if data.get(field):
                data[field] = add_quotes(data[field])

        return data
    except Exception as e:
        logger.error(f"AI extract error: {e}")
        return {"error": str(e)}


# ─── Скачивание шаблона ──────────────────────────────────────────────────────
def download_template(url: str) -> io.BytesIO | None:
    if url in _template_cache:
        return io.BytesIO(_template_cache[url])
    try:
        r = requests.get(url, timeout=30)
        if r.status_code == 200:
            _template_cache[url] = r.content
            return io.BytesIO(r.content)
    except Exception as e:
        logger.error(f"Template download error: {e}")
    return None


# ─── Заполнение шаблона ──────────────────────────────────────────────────────
def fill_contract(template_stream: io.BytesIO, data: dict, manager: dict, contract_key: str, contract_number: str) -> str:
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        raise RuntimeError("Установи docxtpl: pip install docxtpl")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp.write(template_stream.getvalue())
        tmp_path = tmp.name

    try:
        doc = DocxTemplate(tmp_path)
        dates = get_current_date()
        context = {
            **dates,
            'Дата': dates['date_full'],  # {{ Дата }} → «28» мая 2026 г.
            'date': dates['date_digital'],  # {{ date }} → 28.05.2026
            'company_name':          data.get('company_name', '___________'),
            'company_name_short':    data.get('company_name_short', '___________'),
            'company_name_genitive': data.get('company_name_genitive', '___________'),
            'inn':                   data.get('inn', '___________'),
            'kpp':                   data.get('kpp', '___________'),
            'ogrn':                  data.get('ogrn', '___________'),
            'ogrnip':                data.get('ogrnip', '___________'),
            'ceo_name':              data.get('ceo_name_nominative', '___________'),
            'ceo_name_short':        data.get('ceo_name_nominative_short', '___________'),
            'ceo_name_genitive':     data.get('ceo_name_genitive', '___________'),
            'ceo_name_genitive_short': data.get('ceo_name_genitive_short', '___________'),
            'ceo_position':          data.get('ceo_position_nominative', 'Генеральный директор'),
            'ceo_position_genitive': data.get('ceo_position_genitive', '___________'),
            'is_individual_entrepreneur': data.get('is_individual_entrepreneur', False),
            'is_company':            not data.get('is_individual_entrepreneur', False),
            'legal_basis':           data.get('legal_basis', 'Устава'),
            'legal_basis_full':      data.get('legal_basis_full', 'на основании Устава'),
            'acting_genitive':       get_acting_form(data.get('ceo_name_nominative', '')),
            'address':               data.get('address', '___________'),
            'phone':                 data.get('phone') or None,
            'email':                 data.get('email') or None,
            'bik':                   data.get('bik', '___________'),
            'account':               data.get('account', '___________'),
            'korr_account':          data.get('korr_account', '___________'),
            'bank_name':             data.get('bank_name', '___________'),
            'contract_number':       contract_number,
            'city':                  'г. ___________',
            'manager_name':          manager.get('name', ''),
            'manager_position':      manager.get('position', ''),
            'manager_phone':         manager.get('phone', ''),
            'manager_email':         manager.get('email', ''),
            'logist_code':           manager.get('prefix', ''),
        }

        # Определяем пол для правильного склонения
        gender = get_gender(data.get('ceo_name_nominative', ''))
        is_ip = data.get('is_individual_entrepreneur', False)
        context['acting_form'] = 'действующей' if gender == 'f' else 'действующего'
        context['named_form'] = 'именуемая' if (gender == 'f' and is_ip) else 'именуемое'
        doc.render(context)
        out_path = tempfile.NamedTemporaryFile(suffix='.docx', delete=False).name
        doc.save(out_path)
        return out_path
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ─── Хендлеры ────────────────────────────────────────────────────────────────

async def start_contract_flow(message: Message):
    """Вызывается из text.py когда детектируется запрос на договор."""
    user_id = message.from_user.id
    manager = get_manager(user_id)
    CONTRACT_STATE[user_id] = {'step': 'select'}

    name_part = f", {manager['name']}" if manager.get('name') else ""
    await message.answer(
        f"📋 Коллега{name_part}, выбери тип договора:",
        reply_markup=get_contract_keyboard()
    )


@router.message(F.text.func(lambda t: t in [c['button'] for c in CONTRACTS.values()] + ["❌ Отмена"]))
async def handle_contract_button(message: Message):
    """Обрабатывает только нажатия кнопок договоров."""
    user_id = message.from_user.id
    state = CONTRACT_STATE.get(user_id)
    if not state:
        return

    text = message.text.strip()

    # Отмена
    if text == "❌ Отмена":
        CONTRACT_STATE.pop(user_id, None)
        await message.answer(
            "Отменено. Чем могу помочь?",
            reply_markup=ReplyKeyboardRemove()
        )
        return

    # Шаг 1: выбор типа договора
    if state['step'] == 'select':
        selected_key = None
        for key, contract in CONTRACTS.items():
            if contract['button'] == text:
                selected_key = key
                break

        if not selected_key:
            await message.answer(
                "Используй кнопки ниже 👇",
                reply_markup=get_contract_keyboard()
            )
            return

        CONTRACT_STATE[user_id] = {'step': 'wait_card', 'contract_type': selected_key}
        await message.answer(
            f"✅ Выбран: <b>{CONTRACTS[selected_key]['name']}</b>\n\n"
            "📎 Пришли карточку компании — текст, PDF, DOCX или фото.\n"
            "AI извлечёт реквизиты и подставит в договор.",
            parse_mode="HTML",
            reply_markup=ReplyKeyboardRemove()
        )



# Текстовые карточки компании обрабатываются через text.py → _process_card_text_external
async def handle_card_text_if_in_state(message: Message):
    """Вызывается из text.py когда пользователь в режиме wait_card."""
    user_id = message.from_user.id
    state = CONTRACT_STATE.get(user_id)
    if not state or state.get('step') != 'wait_card':
        return False
    text = message.text.strip()
    if len(text) < 30:
        await message.answer("Текст слишком короткий. Пришли карточку компании.")
        return True
    await _process_card_text(message, text)
    return True


@router.message(F.document | F.photo)
async def handle_contract_file(message: Message):
    user_id = message.from_user.id
    state = CONTRACT_STATE.get(user_id)
    if not state or state.get('step') != 'wait_card':
        return  # Не в режиме договора — пропускаем совсем

    status = await message.answer("⏳ Обрабатываю файл...")
    tmp_path = None
    try:
        from bot_instance import bot
        if message.document:
            tg_doc = message.document
            file_obj = await bot.get_file(tg_doc.file_id)
            ext = os.path.splitext(tg_doc.file_name or 'file')[1].lower()
        else:
            tg_photo = message.photo[-1]
            file_obj = await bot.get_file(tg_photo.file_id)
            ext = '.jpg'

        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp_path = tmp.name
            await bot.download_file(file_obj.file_path, destination=tmp_path)

        await status.edit_text("⏳ Извлекаю текст из файла...")
        raw_text = await asyncio.to_thread(extract_text_from_file, tmp_path)

        if not raw_text or len(raw_text.strip()) < 20:
            await status.edit_text("❌ Не удалось извлечь текст. Попробуй другой формат или отправь текстом.")
            return

        await status.delete()
        await _process_card_text(message, raw_text)

    except Exception as e:
        logger.error(f"Contract file error: {e}")
        await status.edit_text(f"❌ Ошибка: {str(e)[:150]}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def _process_card_text(message: Message, raw_text: str):
    user_id = message.from_user.id
    state = CONTRACT_STATE.get(user_id, {})
    contract_key = state.get('contract_type')
    if not contract_key:
        return

    status = await message.answer("🧠 AI анализирует реквизиты...")
    try:
        company_data = await extract_company_data(raw_text)
        if company_data.get('error'):
            await status.edit_text(f"❌ Ошибка AI: {company_data['error'][:150]}")
            return

        manager = get_manager(user_id)
        contract_number = await get_contract_number(manager.get('prefix', ''), contract_key)
        company_data['contract_number'] = contract_number

        await status.edit_text("✍️ Заполняю договор...")

        template_stream = await asyncio.to_thread(download_template, CONTRACTS[contract_key]['template_url'])
        if not template_stream:
            await status.edit_text("❌ Не удалось загрузить шаблон с Google Drive.")
            return

        filled_path = await asyncio.to_thread(
            fill_contract, template_stream, company_data, manager, contract_key, contract_number
        )

        short_name = company_data.get('company_name_short', 'client')
        inn = company_data.get('inn', '')
        filename = f"{contract_key}_{inn or short_name}.docx"

        from aiogram.types import FSInputFile
        doc_file = FSInputFile(filled_path, filename=filename)
        await message.answer_document(
            document=doc_file,
            caption=(
                f"✅ <b>{CONTRACTS[contract_key]['name']}</b> готов!\n"
                f"📋 Номер: <code>{contract_number}</code>\n"
                f"🏢 {company_data.get('company_name_short', '')}"
            ),
            parse_mode="HTML"
        )
        os.unlink(filled_path)
        await status.delete()

        CONTRACT_STATE.pop(user_id, None)
        await message.answer(
            "🎉 Готово! Для нового договора напиши «нужен договор».",
            reply_markup=ReplyKeyboardRemove()
        )

    except Exception as e:
        logger.error(f"Contract process error: {e}")
        await status.edit_text(f"❌ Ошибка: {str(e)[:200]}")
        CONTRACT_STATE.pop(user_id, None)
