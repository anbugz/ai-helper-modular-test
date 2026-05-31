"""
services/ocr.py — единый модуль извлечения текста из файлов.

Поддерживаемые форматы:
  .pdf              → pdfplumber → PyPDF2 (fallback)
  .docx             → python-docx
  .xlsx / .xls      → openpyxl
  .txt              → plain text
  .jpg/.jpeg/.png/.webp/.bmp/.tiff → DeepSeek Vision (base64)

Использование:
    from services.ocr import extract_text

    text = await extract_text(file_path)   # для async контекста
    text = extract_text_sync(file_path)    # для sync контекста
"""

import os
import base64
import asyncio
import logging
import httpx
from config import DEEPSEEK_API_KEY, logger

# ─── Константы ────────────────────────────────────────────────────────────────

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".tif"}
TEXT_EXTENSIONS   = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt"}

_TIMEOUT = httpx.Timeout(connect=10.0, read=90.0, write=10.0, pool=5.0)

# ─── Синхронные извлекатели ───────────────────────────────────────────────────

def _extract_pdf(path: str) -> str:
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        if text.strip():
            return text
    except Exception as e:
        logger.debug(f"pdfplumber failed: {e}")

    # Fallback: PyPDF2
    try:
        import PyPDF2
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        logger.debug(f"PyPDF2 failed: {e}")

    return text


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"DOCX extract error: {e}")
        return ""


def _extract_xlsx(path: str) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"XLSX extract error: {e}")
        return ""


def _extract_txt(path: str) -> str:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return ""


def extract_text_sync(path: str) -> str:
    """Синхронное извлечение текста — только для текстовых форматов."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(path)
    elif ext in (".xlsx", ".xls"):
        return _extract_xlsx(path)
    elif ext == ".txt":
        return _extract_txt(path)
    else:
        return ""  # изображения — только через async extract_text


# ─── DeepSeek Vision ──────────────────────────────────────────────────────────

async def _extract_image_via_vision(path: str) -> str:
    """
    Отправляет изображение в DeepSeek Vision и извлекает весь текст.
    Использует deepseek-chat с vision capabilities (multimodal).
    """
    if not DEEPSEEK_API_KEY:
        return ""

    ext = os.path.splitext(path)[1].lower().lstrip(".")
    mime_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "png": "image/png", "webp": "image/webp",
        "bmp": "image/bmp", "tiff": "image/tiff", "tif": "image/tiff",
    }
    mime = mime_map.get(ext, "image/jpeg")

    try:
        with open(path, "rb") as f:
            img_b64 = base64.b64encode(f.read()).decode()
    except Exception as e:
        logger.error(f"Image read error: {e}")
        return ""

    payload = {
        "model": "deepseek-chat",
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime};base64,{img_b64}"
                        }
                    },
                    {
                        "type": "text",
                        "text": (
                            "Извлеки весь текст с изображения дословно. "
                            "Сохрани структуру: таблицы — через |, строки — через новую строку. "
                            "Не добавляй комментариев, только текст документа."
                        )
                    }
                ]
            }
        ],
        "max_tokens": 4000,
    }
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json",
    }

    try:
        limits = httpx.Limits(max_keepalive_connections=0, max_connections=10)
        async with httpx.AsyncClient(timeout=_TIMEOUT, limits=limits) as client:
            resp = await client.post(
                "https://api.deepseek.com/v1/chat/completions",
                json=payload,
                headers=headers,
            )
        if resp.status_code != 200:
            logger.error(f"DeepSeek Vision HTTP {resp.status_code}: {resp.text[:300]}")
            return ""
        data = resp.json()
        return data["choices"][0]["message"].get("content", "")
    except httpx.ReadTimeout:
        logger.error("DeepSeek Vision read timeout")
        return ""
    except Exception as e:
        logger.error(f"DeepSeek Vision error: {e}")
        return ""


# ─── Основная функция ─────────────────────────────────────────────────────────

async def extract_text(path: str) -> str:
    """
    Главная функция — извлекает текст из любого поддерживаемого файла.

    Для текстовых форматов — синхронно в thread pool.
    Для изображений — через DeepSeek Vision.

    Returns:
        Извлечённый текст или "" если не удалось.
    """
    ext = os.path.splitext(path)[1].lower()

    if ext in IMAGE_EXTENSIONS:
        logger.info(f"[OCR] image → DeepSeek Vision: {os.path.basename(path)}")
        return await _extract_image_via_vision(path)

    if ext in TEXT_EXTENSIONS:
        logger.info(f"[OCR] text extract: {os.path.basename(path)}")
        text = await asyncio.to_thread(extract_text_sync, path)
        # Если PDF оказался сканом (пустой текст) — пробуем Vision
        if not text.strip() and ext == ".pdf":
            logger.info(f"[OCR] PDF has no text layer, trying Vision: {os.path.basename(path)}")
            # Конвертируем первые страницы PDF в изображение если возможно
            text = await _pdf_via_vision(path)
        return text

    logger.warning(f"[OCR] unsupported format: {ext}")
    return ""


async def _pdf_via_vision(path: str) -> str:
    """
    Fallback для PDF-сканов: конвертирует страницы в PNG и отправляет в Vision.
    Требует pdf2image + poppler.
    """
    try:
        from pdf2image import convert_from_path
        images = await asyncio.to_thread(
            convert_from_path, path, dpi=200, first_page=1, last_page=3
        )
    except Exception as e:
        logger.debug(f"pdf2image not available: {e}")
        return ""

    import tempfile
    texts = []
    for img in images:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name
        try:
            await asyncio.to_thread(img.save, tmp_path, "PNG")
            t = await _extract_image_via_vision(tmp_path)
            if t.strip():
                texts.append(t)
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    return "\n\n".join(texts)
